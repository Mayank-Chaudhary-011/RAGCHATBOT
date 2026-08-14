import os 
import logging
from docx import Document as DocxDocument
from langchain_core.documents import Document as LangChainDoc

logger = logging.getLogger(__name__)


def load_docx_folder(folder_path:str = "docs") -> list[LangChainDoc]:
    """Load all .docx 3GPP spec files from folder as LangChain
        Documents"""
    documents:list[LangChainDoc] = []

    for filename in os.listdir(folder_path):
        if not filename.lower().endswith(".docx"):
            continue

        path = os.path.join(folder_path,filename)
        try:
            doc = DocxDocument(path)
        except Exception as e:
            logger.warning(f"Skiping {filename}: {e}")
            continue


        # Cleaning Data
        
        para_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        table_text = "\n".join(
            " | ".join(c.text.strip() for c in row.cells)
            for table in doc.tables
            for row in table.rows
            if any(c.text.strip() for c in row.cells)
        )

        text ="\n".join(filter(None, [para_text , table_text]))
        text = text.replace("\xa0"," ")

        spec_id = os.path.splitext(filename)[0].split("-")[0]

        documents.append(LangChainDoc(
            page_content=text,
            metadata={"source":filename,"spec_id":spec_id,"file_path":path}

        ))
    print(f"\nTotal documents loaded:{len(documents)}")
    return documents